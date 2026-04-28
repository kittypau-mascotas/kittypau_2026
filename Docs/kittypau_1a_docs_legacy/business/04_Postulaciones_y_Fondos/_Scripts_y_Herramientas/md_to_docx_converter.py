import os
from docx import Document
from docx.shared import Inches

def convert_md_to_docx(md_file_path, docx_output_path):
    try:
        document = Document()
        
        with open(md_file_path, 'r', encoding='utf-8') as md_file:
            for line in md_file:
                line = line.strip()
                if not line:
                    document.add_paragraph('') # Add empty paragraph for blank lines
                    continue

                # Basic Markdown parsing
                if line.startswith('### '):
                    document.add_heading(line[4:], level=3)
                elif line.startswith('## '):
                    document.add_heading(line[3:], level=2)
                elif line.startswith('# '):
                    document.add_heading(line[2:], level=1)
                elif line.startswith('* ') or line.startswith('- '):
                    document.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith('**') and line.endswith('**'):
                    document.add_paragraph(line.strip('**')).bold = True
                elif line.startswith('*') and line.endswith('*'):
                    document.add_paragraph(line.strip('*')).italic = True
                else:
                    document.add_paragraph(line)
        
        document.save(docx_output_path)
        print(f"Convertido exitosamente {md_file_path} a {docx_output_path}")
        return True
    except Exception as e:
        print(f"Error al convertir Markdown {md_file_path} a DOCX: {e}")
        return False

if __name__ == "__main__":
    # Define la ruta al archivo Markdown de la plantilla
    md_template_path = r"D:\Escritorio\Proyectos\Kittypau\Kittypau_1a\docs\business\04_Postulaciones_y_Fondos\01_Capital_Semilla_2025\Anexos\Carta_Compromiso_Ignacio_Farias.md"
    
    # Define la ruta de salida para el archivo DOCX
    # Se colocará en el mismo directorio que la plantilla Markdown
    docx_output_path = os.path.splitext(md_template_path)[0] + ".docx"

    print(f"Intentando convertir {md_template_path} a DOCX...")
    convert_md_to_docx(md_template_path, docx_output_path)
    print("Proceso de conversión finalizado.")
