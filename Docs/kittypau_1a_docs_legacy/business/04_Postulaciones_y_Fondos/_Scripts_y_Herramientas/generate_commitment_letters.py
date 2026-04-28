import os
import re
from docx import Document
import pandas as pd

def generate_commitment_letters(template_md_path, data_xlsx_path, output_dir):
    try:
        # Ensure output directory exists
        os.makedirs(output_dir, exist_ok=True)

        # Read the Markdown template content
        with open(template_md_path, 'r', encoding='utf-8') as f:
            template_content = f.read()

        # Read data from XLSX using pandas
        df = pd.read_excel(data_xlsx_path)

        # Extract placeholders from the template
        placeholders = re.findall(r'[(.*?)]', template_content)
        
        # Convert Markdown content to DOCX (basic conversion)
        def create_docx_from_markdown_content(markdown_text, docx_document):
            for line in markdown_text.split('\n'):
                line = line.strip()
                if not line:
                    docx_document.add_paragraph('')
                    continue

                if line.startswith('### '):
                    docx_document.add_heading(line[4:], level=3)
                elif line.startswith('## '):
                    docx_document.add_heading(line[3:], level=2)
                elif line.startswith('# '):
                    docx_document.add_heading(line[2:], level=1)
                elif line.startswith('* ') or line.startswith('- '):
                    docx_document.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith('**') and line.endswith('**'):
                    docx_document.add_paragraph(line.strip('**')).bold = True
                elif line.startswith('*') and line.endswith('*'):
                    docx_document.add_paragraph(line.strip('*')).italic = True
                else:
                    docx_document.add_paragraph(line)

        # Iterate over each row (participant) in the DataFrame
        for index, row in df.iterrows():
            filled_content = template_content
            output_filename = f"Carta_Compromiso_{row['Nombre Completo del Participante'].replace(' ', '_')}.docx"
            output_docx_path = os.path.join(output_dir, output_filename)
            
            # Replace each placeholder with data from the current row
            for placeholder in placeholders:
                if placeholder in row:
                    filled_content = filled_content.replace(f'[{placeholder}]', str(row[placeholder]))
                else:
                    print(f"Advertencia: El marcador de posición '{placeholder}' no se encontró en el archivo XLSX para el participante {row['Nombre Completo del Participante']}.")
                    filled_content = filled_content.replace(f'[{placeholder}]', '[FALTA DATO]')
            
            # Create a new DOCX document for each participant
            doc = Document()
            create_docx_from_markdown_content(filled_content, doc)
            doc.save(output_docx_path)
            print(f"Generada carta para {row['Nombre Completo del Participante']} en {output_docx_path}")

        print("Proceso de generación de cartas finalizado.")

    except FileNotFoundError as e:
        print(f"Error: Archivo no encontrado - {e}")
        print("Asegúrate de que las rutas a la plantilla Markdown y al archivo XLSX sean correctas.")
    except KeyError as e:
        print(f"Error: Columna '{e}' no encontrada en el archivo XLSX.")
        print("Asegúrate de que los nombres de las columnas en tu XLSX coincidan exactamente con los marcadores de posición en la plantilla Markdown (sin los corchetes).")
    except Exception as e:
        print(f"Ocurrió un error inesperado: {e}")

if __name__ == "__main__":
    # Define paths
    template_md_path = "D:\\Escritorio\\Proyectos\\Kittypau\\Kittypau_1a\\docs\\business\\04_Postulaciones_y_Fondos\\01_Capital_Semilla_2025\\Anexos\\Carta_Compromiso_Piloto_Kittypau_Template.md"
    # IMPORTANT: Replace with the actual path to your XLSX data file
    data_xlsx_path = "D:\\Escritorio\\Proyectos\\Kittypau\\Kittypau_1a\\docs\\business\\04_Postulaciones_y_Fondos\\datos_participantes_piloto.xlsx" # Example path
    output_letters_dir = "D:\\Escritorio\\Proyectos\\Kittypau\\Kittypau_1a\\docs\\business\\04_Postulaciones_y_Fondos\\01_Capital_Semilla_2025\\Cartas_Generadas"

    print("Iniciando la generación de cartas de compromiso...")
    generate_commitment_letters(template_md_path, data_xlsx_path, output_letters_dir)

